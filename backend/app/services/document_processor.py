import re
import logging
from typing import Optional, Dict, Any
from pathlib import Path
from io import BytesIO

# PDF processing
import PyPDF2
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdfminer.layout import LAParams

# DOCX processing
from docx import Document

from app.schemas.file_upload import FileType
from app.core.exceptions import DocumentExtractionError, FileProcessingError
from app.core.error_handler import error_handler, RetryHandler
from app.core.logging_config import performance_logger

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for extracting and processing text from various document formats"""
    
    def __init__(self):
        self.supported_formats = {FileType.PDF, FileType.DOCX, FileType.TXT}
    
    async def extract_text(self, file_path: Path, file_type: FileType) -> str:
        """
        Extract text from a document file with comprehensive error handling.
        
        Args:
            file_path: Path to the document file
            file_type: Type of the document (PDF, DOCX, TXT)
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentExtractionError: If extraction fails
            FileProcessingError: If file access fails
        """
        with performance_logger.time_operation(
            "document_extraction",
            file_type=file_type.value,
            file_size=file_path.stat().st_size if file_path.exists() else 0
        ):
            if not file_path.exists():
                raise await error_handler.handle_file_processing_error(
                    FileNotFoundError(f"File not found: {file_path}"),
                    file_path=str(file_path),
                    operation="file_access"
                )
            
            if file_type not in self.supported_formats:
                raise DocumentExtractionError(
                    f"Unsupported file type: {file_type}",
                    file_type=file_type.value
                )
            
            try:
                # Use retry logic for extraction
                return await RetryHandler.retry_with_backoff(
                    self._extract_text_with_fallback,
                    max_retries=2,
                    base_delay=1.0,
                    exceptions=(DocumentExtractionError, IOError, OSError),
                    file_path=file_path,
                    file_type=file_type
                )
            
            except Exception as e:
                logger.error(f"Failed to extract text from {file_path}: {str(e)}")
                raise await error_handler.handle_document_extraction_error(
                    e,
                    file_type=file_type.value,
                    operation="text_extraction"
                )
    
    async def _extract_text_with_fallback(self, file_path: Path, file_type: FileType) -> str:
        """Extract text with fallback methods."""
        if file_type == FileType.PDF:
            return await self._extract_pdf_text(file_path)
        elif file_type == FileType.DOCX:
            return await self._extract_docx_text(file_path)
        elif file_type == FileType.TXT:
            return await self._extract_txt_text(file_path)
        else:
            raise DocumentExtractionError(f"No extraction method for file type: {file_type}")
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file using multiple methods for better reliability"""
        text = ""
        
        # Try PyPDF2 first (faster)
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If PyPDF2 extracted reasonable amount of text, use it
            if len(text.strip()) > 50:
                return text.strip()
        
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed for {file_path}: {str(e)}")
        
        # Fallback to pdfminer for better text extraction
        try:
            laparams = LAParams(
                line_margin=0.5,
                word_margin=0.1,
                char_margin=2.0,
                boxes_flow=0.5,
                all_texts=False
            )
            text = pdfminer_extract_text(str(file_path), laparams=laparams)
            return text.strip() if text else ""
        
        except Exception as e:
            logger.error(f"pdfminer extraction failed for {file_path}: {str(e)}")
            raise DocumentExtractionError(f"PDF text extraction failed: {str(e)}")
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            return "\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
            raise DocumentExtractionError(f"DOCX text extraction failed: {str(e)}")
    
    async def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue
            
            raise DocumentExtractionError("Could not decode text file with any supported encoding")
        
        except Exception as e:
            logger.error(f"TXT extraction failed for {file_path}: {str(e)}")
            raise DocumentExtractionError(f"TXT text extraction failed: {str(e)}")
    
    async def normalize_text(self, raw_text: str) -> str:
        """
        Normalize and clean extracted text.
        
        Args:
            raw_text: Raw extracted text
            
        Returns:
            Normalized and cleaned text
        """
        if not raw_text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', raw_text)
        
        # Remove special characters that might interfere with processing
        text = re.sub(r'[^\w\s\-.,;:()\[\]{}@#$%&*+=<>?/\\|`~\'"]', ' ', text)
        
        # Remove multiple consecutive punctuation marks
        text = re.sub(r'([.,:;!?]){2,}', r'\1', text)
        
        # Normalize line breaks and spacing
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Preserve paragraph breaks
        text = re.sub(r'\n(?!\n)', ' ', text)    # Convert single line breaks to spaces
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    async def get_document_metadata(self, file_path: Path, file_type: FileType) -> Dict[str, Any]:
        """
        Extract metadata from document.
        
        Args:
            file_path: Path to the document file
            file_type: Type of the document
            
        Returns:
            Dictionary containing document metadata
        """
        metadata = {
            "file_path": str(file_path),
            "file_type": file_type.value,
            "file_size": file_path.stat().st_size if file_path.exists() else 0,
        }
        
        try:
            if file_type == FileType.PDF:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata.update({
                        "page_count": len(pdf_reader.pages),
                        "pdf_info": pdf_reader.metadata if pdf_reader.metadata else {}
                    })
            
            elif file_type == FileType.DOCX:
                doc = Document(file_path)
                metadata.update({
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "core_properties": {
                        "title": doc.core_properties.title,
                        "author": doc.core_properties.author,
                        "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                        "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None,
                    }
                })
        
        except Exception as e:
            logger.warning(f"Failed to extract metadata from {file_path}: {str(e)}")
            metadata["metadata_error"] = str(e)
        
        return metadata
    
    def validate_extracted_text(self, text: str, min_length: int = 10) -> bool:
        """
        Validate that extracted text is meaningful.
        
        Args:
            text: Extracted text to validate
            min_length: Minimum required text length
            
        Returns:
            True if text appears valid, False otherwise
        """
        if not text or len(text.strip()) < min_length:
            return False
        
        # Check if text contains mostly readable characters
        readable_chars = sum(1 for c in text if c.isalnum() or c.isspace() or c in '.,;:!?-()[]{}')
        total_chars = len(text)
        
        if total_chars == 0:
            return False
        
        readable_ratio = readable_chars / total_chars
        return readable_ratio > 0.7  # At least 70% readable characters


# Global document processor instance
document_processor = DocumentProcessor()